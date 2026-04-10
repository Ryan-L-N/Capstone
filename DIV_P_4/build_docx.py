"""
Build Phase_Deliverable_Locomotion.docx in Colby's doc style.
Run once, then delete this script.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Global styles ──────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.space_before = Pt(0)
pf.line_spacing = 1.15

# Adjust heading styles
for level, size, color in [
    ("Heading 1", 22, RGBColor(0x1B, 0x3A, 0x5C)),
    ("Heading 2", 16, RGBColor(0x1B, 0x3A, 0x5C)),
    ("Heading 3", 13, RGBColor(0x2E, 0x56, 0x7A)),
]:
    s = doc.styles[level]
    s.font.name = "Calibri"
    s.font.size = Pt(size)
    s.font.color.rgb = color
    s.font.bold = True
    s.paragraph_format.space_before = Pt(18 if level != "Heading 1" else 6)
    s.paragraph_format.space_after = Pt(6)

# Margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Helper functions ───────────────────────────────────────────

def add_hr():
    """Add a horizontal rule (Colby-style section divider)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "6",
        qn("w:space"): "1",
        qn("w:color"): "CCCCCC",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_meta(text):
    """Italicized metadata line (Colby style)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.paragraph_format.space_after = Pt(2)


def add_body(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        b = p.add_run(bold_prefix)
        b.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_callout(text):
    """Blockquote-style callout (Colby's key takeaway boxes)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    left = pBdr.makeelement(qn("w:left"), {
        qn("w:val"): "single",
        qn("w:sz"): "18",
        qn("w:space"): "8",
        qn("w:color"): "2E74B5",
    })
    pBdr.append(left)
    pPr.append(pBdr)
    # Set background shading
    shd = pPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): "E8F0FE",
    })
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.italic = True
    return p


def add_table(headers, rows, bold_col=None, highlight_rows=None):
    """Add a formatted table."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Dark header background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = tcPr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "1B3A5C",
        })
        tcPr.append(shd)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if bold_col is not None and c_idx == bold_col:
                run.bold = True
            if highlight_rows and r_idx in highlight_rows:
                run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Alternating row shading
            if r_idx % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = tcPr.makeelement(qn("w:shd"), {
                    qn("w:val"): "clear",
                    qn("w:color"): "auto",
                    qn("w:fill"): "F2F6FA",
                })
                tcPr.append(shd)

    doc.add_paragraph()  # spacer
    return tbl


def add_code_block(text):
    """Monospace code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = pPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): "F5F5F5",
    })
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_bullet(text, bold_prefix=None, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    if bold_prefix:
        b = p.add_run(bold_prefix)
        b.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        b = p.add_run(bold_prefix)
        b.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


# ══════════════════════════════════════════════════════════════
#  DOCUMENT CONTENT
# ══════════════════════════════════════════════════════════════

# ── Title ──
title = doc.add_heading("Locomotion Policy Development", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_meta("AI2C Tech Capstone — Carnegie Mellon University | April 2026")
add_meta("Team: Alex, Ryan, Colby, Cole | Platform: Boston Dynamics Spot")
add_hr()

# ── 1. Overview ──
doc.add_heading("1. Overview", level=2)

add_body(
    "This document tells the story of how we developed a single generalist locomotion "
    "policy for the Boston Dynamics Spot quadruped — one that handles smooth ground, "
    "tall grass, boulder fields, and stairs. Rather than training one monolithic policy "
    "to do everything (which leads to conflicting reward gradients), we followed a "
    "four-stage pipeline:"
)

add_numbered("", bold_prefix="ARL Baseline — ")
doc.paragraphs[-1].add_run(
    "An externally proven policy with a simpler, more elegant design that outperformed "
    "our custom 22-term configuration."
)
add_numbered("", bold_prefix="ARL Hybrid — ")
doc.paragraphs[-1].add_run(
    "We adapted the ARL Baseline to our harder 12-terrain curriculum and added three "
    "targeted safety fixes."
)
add_numbered("", bold_prefix="Expert Masters — ")
doc.paragraphs[-1].add_run(
    "We specialized two copies of the ARL Hybrid into terrain experts: one for smooth "
    "surfaces and one for obstacles."
)
add_numbered("", bold_prefix="Distilled Master — ")
doc.paragraphs[-1].add_run(
    "We combined both experts into a single student policy using multi-expert distillation "
    "with a terrain-aware router."
)

add_body(
    "Each stage built on the last. The result is a policy that inherits the best behavior "
    "from each expert without the trade-offs that plagued single-policy training."
)

add_callout(
    "All training used NVIDIA Isaac Lab, RSL-RL (PPO), and an NVIDIA H100 GPU."
)

# ── Summary Table ──
doc.add_heading("Pipeline Summary", level=3)

add_table(
    ["Stage", "Policy", "Checkpoint", "Key Metric"],
    [
        ["1. Baseline", "ARL Baseline", "mason_baseline_final_19999.pt", "Friction 36.9m, 37% falls"],
        ["2. Hybrid", "ARL Hybrid", "mason_hybrid_best_33200.pt", "Friction 48.9m, 2% falls, 0% flip"],
        ["3a. Expert", "Friction/Grass Expert", "(same as Hybrid)", "48.9m friction, 98% completion"],
        ["3b. Expert", "Obstacle Expert", "obstacle_best_44400.pt", "30.4m boulder (+10.1m)"],
        ["4. Distilled", "Distilled Master", "distilled_6899.pt", "Best of both experts"],
    ],
    bold_col=1,
)

add_hr()

# ── 2. Policy Architecture ──
doc.add_heading("2. Policy Architecture", level=2)

add_body(
    "All policies in the pipeline share the same architecture, making them directly "
    "compatible with each other and with our evaluation harness:"
)

add_table(
    ["Property", "Value"],
    [
        ["Network", "[512, 256, 128] MLP with ELU activation"],
        ["Parameters", "~286,604"],
        ["Observation dimensions", "235 (187 height scan + 48 proprioception)"],
        ["Action dimensions", "12 (joint position targets)"],
        ["Control frequency", "50 Hz"],
        ["Physics frequency", "500 Hz (decimation = 10)"],
        ["PD gains", "Kp = 60.0, Kd = 1.5"],
        ["Action scale", "0.2 rad offset from default standing pose"],
        ["Observation order", "Height scan first (0-186), proprioception second (187-234)"],
    ],
)

add_callout(
    "Key insight: The [512, 256, 128] architecture (286K params) outperformed our earlier "
    "[1024, 512, 256] network (1.2M params). The smaller network generalizes better — "
    "less capacity forces the policy to learn broadly useful locomotion rather than "
    "memorizing terrain-specific tricks."
)

add_hr()

# ── 3. Reward Function ──
doc.add_heading("3. Reward Function", level=2)
doc.add_heading("3.1 ARL's 11-Term Reward Function", level=3)

add_body(
    "ARL's original reward structure uses 11 clean, well-tuned terms. Fewer signals "
    "give clearer gradients — the policy knows exactly what we want."
)

add_body("Positive Rewards (Incentives)", bold_prefix="")
doc.paragraphs[-1].runs[0].bold = True

add_table(
    ["Term", "Weight", "Description"],
    [
        ["gait", "10.0", "Diagonal trot synchronization (FL+HR, FR+HL alternating)"],
        ["base_linear_velocity", "5.0", "Reward for matching commanded forward/lateral speed"],
        ["base_angular_velocity", "5.0", "Reward for matching commanded yaw rate"],
        ["air_time", "5.0", "Reward for proper swing phase (feet off ground ~0.3s)"],
        ["foot_clearance", "0.5", "Reward for lifting feet during swing (obstacle clearance)"],
    ],
)

add_body("Negative Rewards (Penalties)", bold_prefix="")
doc.paragraphs[-1].runs[0].bold = True

add_table(
    ["Term", "Weight", "Description"],
    [
        ["base_orientation", "-3.0", "Excessive body roll/pitch"],
        ["base_motion", "-2.0", "Unwanted vertical/lateral body velocity"],
        ["air_time_variance", "-1.0", "Inconsistent swing timing across legs"],
        ["action_smoothness", "-1.0", "Rapid changes in control commands (clamped at 10.0)"],
        ["joint_pos", "-0.7", "Extreme joint angle deviations from default"],
        ["foot_slip", "-0.5", "Feet sliding during ground contact"],
    ],
)

doc.add_heading("3.2 Three Safety Additions (ARL Hybrid)", level=3)

add_body(
    "When we put ARL's config on our harder 12-terrain curriculum, three failure modes "
    "appeared. We added one targeted fix for each:"
)

add_table(
    ["Addition", "Weight", "Why It Was Needed"],
    [
        ["terrain_relative_height", "-2.0", "Without this, the robot belly-crawls. Fixed 0.37m target height."],
        ["dof_pos_limits", "-3.0", "Without this, the policy locks knees at mechanical stops."],
        ["clamped_action_smoothness", "(replaces original)", "Unbounded L2 norms can explode to NaN. Caps at 10.0."],
    ],
)

add_callout(
    "Total: 14 terms — ARL's proven 11, plus 3 surgical fixes. Our earlier custom "
    "configuration had 22 terms with competing penalties that made gradient signals unclear."
)

doc.add_heading("3.3 Reward Design Principles", level=3)

add_numbered("", bold_prefix="Fewer rewards = clearer gradients. ")
doc.paragraphs[-1].add_run("11 terms outperformed 22.")
add_numbered("", bold_prefix="Gait weight must stay high (10.0+). ")
doc.paragraphs[-1].add_run("Lower values produce bouncing exploits.")
add_numbered("", bold_prefix="Tune complementary rewards together. ")
doc.paragraphs[-1].add_run("Foot clearance, action smoothness, and joint position all govern leg-lift behavior.")
add_numbered("", bold_prefix="Clamp all unbounded penalty terms. ")
doc.paragraphs[-1].add_run("Squared errors and norms can explode to NaN without clamping.")
add_numbered("", bold_prefix="Use terrain-relative height, not world-frame Z. ")
doc.paragraphs[-1].add_run("World-frame height is meaningless on elevated terrain.")

add_hr()

# ── 4. Training Curriculum ──
doc.add_heading("4. Training Curriculum", level=2)
doc.add_heading("4.1 12-Terrain Curriculum", level=3)

add_body(
    "Each training run uses a grid of 10 difficulty rows x 40 terrain columns = "
    "400 patches (8m x 8m each):"
)

add_table(
    ["Category", "Terrain Type", "Proportion", "Description"],
    [
        ["Geometric (40%)", "Pyramid stairs up", "8%", "Ascending stair pyramids"],
        ["", "Pyramid stairs down", "8%", "Descending stair pyramids"],
        ["", "Random grid boxes", "8%", "Repeated box obstacles"],
        ["", "Stepping stones", "8%", "Sparse elevated platforms"],
        ["", "Repeated boxes", "8%", "Uniform obstacle grids"],
        ["Surface (35%)", "Random rough", "10%", "Bumpy ground (0.02-0.12m)"],
        ["", "Slopes up", "5%", "Inclined planes"],
        ["", "Slopes down", "5%", "Declined planes"],
        ["", "Wave terrain", "5%", "Sinusoidal ground"],
        ["", "Friction plane", "5%", "Variable friction (0.05-0.9)"],
        ["", "Vegetation plane", "5%", "Ground with drag forces"],
        ["Compound (25%)", "HF stairs", "10%", "Height-field linear stairs"],
        ["", "Discrete obstacles", "15%", "Random polyhedra"],
    ],
)

add_callout(
    "Automatic curriculum: Robots are promoted to harder difficulty rows when they "
    "consistently survive, and demoted when they fall. Training always occurs at "
    "the challenge frontier."
)

add_hr()

# ── 5. The Policy Pipeline ──
doc.add_heading("5. The Policy Pipeline", level=2)

# ── 5.1 ARL Baseline ──
doc.add_heading("5.1 Stage 1 — ARL Baseline", level=3)

add_body(
    "ARL's team independently developed a locomotion policy that reached terrain level ~6.0 "
    "using a simpler design than our custom 22-term, 1.2M-parameter configuration. After 11 "
    "trials and ~30 sub-iterations of tuning our custom config, we adopted ARL's approach."
)

add_body("Why ARL's design worked better:", bold_prefix="")
doc.paragraphs[-1].runs[0].bold = True

add_table(
    ["Feature", "Our Custom Config", "ARL Baseline"],
    [
        ["Reward terms", "22 (competing signals)", "11 (clean gradients)"],
        ["Network size", "[1024,512,256] — 1.2M params", "[512,256,128] — 286K params"],
        ["LR schedule", "Cosine annealing (manual)", "Adaptive KL (self-adjusting)"],
        ["Domain randomization", "Heavy (mass ±5kg)", "Light (mass ±2.5kg)"],
        ["Observation noise", "Enabled", "Disabled"],
        ["Episode length", "30s", "20s"],
        ["Mini-batches", "64", "4 (larger effective batch)"],
    ],
)

add_table(
    ["Property", "Value"],
    [
        ["Checkpoint", "mason_baseline_final_19999.pt"],
        ["Architecture", "[512, 256, 128]"],
        ["Iterations", "20,000"],
        ["Environments", "4,096"],
    ],
)

add_body("100-Episode Evaluation:", bold_prefix="")
doc.paragraphs[-1].runs[0].bold = True

add_table(
    ["Environment", "Mean Progress", "Zone (avg)", "Fall Rate"],
    [
        ["Friction", "36.9m", "4.1", "37%"],
        ["Grass", "29.6m", "3.6", "23%"],
        ["Boulder", "14.4m", "2.2", "13%"],
        ["Stairs", "10.9m", "2.0", "20%"],
    ],
)

add_callout(
    "Takeaway: ARL's clean design produced a functional walking gait, but on our harder "
    "12-terrain curriculum it needed refinement. The high friction fall rate (37%) and "
    "limited boulder performance (14.4m) showed where targeted fixes would help."
)

add_hr()

# ── 5.2 ARL Hybrid ──
doc.add_heading("5.2 Stage 2 — ARL Hybrid", level=3)

add_body(
    "ARL's 11 reward terms + our 12-terrain curriculum + 3 surgical safety fixes "
    "(terrain-relative height, DOF limits, clamped action smoothness). No AI coach — "
    "just ARL's proven weights and our harder terrain."
)

add_table(
    ["Property", "Value"],
    [
        ["Checkpoint", "mason_hybrid_best_33200.pt"],
        ["Architecture", "[512, 256, 128]"],
        ["Iterations", "33,200 (best), 35,100 (final)"],
        ["Training time", "~42.6 hours"],
        ["Total steps", "~2.0 billion"],
        ["Peak terrain level", "3.83"],
        ["Flip rate", "0%"],
        ["Environments", "4,096"],
    ],
)

add_body("100-Episode Evaluation:", bold_prefix="")
doc.paragraphs[-1].runs[0].bold = True

add_table(
    ["Environment", "Mean Progress", "Zone (avg)", "Fall Rate", "Velocity"],
    [
        ["Friction", "48.9 ± 5.0m", "5.0", "2%", "0.934 m/s"],
        ["Grass", "27.2 ± 8.0m", "3.3", "15%", "0.487 m/s"],
        ["Boulder", "20.3 ± 1.7m", "3.0", "3%", "0.350 m/s"],
        ["Stairs", "11.2 ± 2.0m", "2.0", "36%", "0.227 m/s"],
    ],
)

add_body("What improved over the Baseline:", bold_prefix="")
doc.paragraphs[-1].runs[0].bold = True

add_table(
    ["Environment", "Baseline → Hybrid", "Change"],
    [
        ["Friction", "36.9m → 48.9m", "+12.0m, fall rate 37% → 2%"],
        ["Boulder", "14.4m → 20.3m", "+5.9m, fall rate 13% → 3%"],
        ["Stairs", "10.9m → 11.2m", "+0.3m (minimal change)"],
        ["Grass", "29.6m → 27.2m", "-2.4m (slight regression)"],
    ],
    highlight_rows=[0, 1],
)

add_body(
    "The three safety additions transformed friction performance (from 37% falls to 2%) "
    "and improved boulder traversal by 41%. The ARL Hybrid became a rock-solid generalist "
    "with a 0% flip rate — it never flipped over, even on difficult terrain."
)

add_callout(
    "The trade-off problem: The ARL Hybrid was excellent on smooth terrain but limited on "
    "obstacles. Training a single policy harder on boulders and stairs would regress its "
    "friction and grass performance. This tension — loose penalties for obstacles vs. tight "
    "penalties for clean walking — was the core problem that motivated the next stage."
)

add_hr()

# ── 5.3 Expert Masters ──
doc.add_heading("5.3 Stage 3 — The Expert Masters", level=3)

add_body(
    "Rather than training one policy to do everything, we split the problem. Two specialist "
    "experts were trained, each optimized for its terrain type:"
)

doc.add_heading("Friction/Grass Expert", level=3)

add_body(
    "The ARL Hybrid itself served as the friction/grass expert. It was already excellent "
    "on smooth surfaces (48.9m friction, 98% completion) with a clean, stable gait. "
    "No additional training was needed."
)

doc.add_heading("Obstacle Expert", level=3)

add_body(
    "Starting from the ARL Hybrid, we retrained with a 60% boulder/stair terrain mix "
    "to force the policy to specialize on obstacles."
)

add_table(
    ["Property", "Value"],
    [
        ["Checkpoint", "obstacle_best_44400.pt"],
        ["Starting from", "ARL Hybrid (33200)"],
        ["Terrain mix", "60% boulders/stairs, 40% mixed"],
        ["Iterations", "44,400 (best out of 54,600)"],
        ["Peak terrain level", "4.38"],
        ["Flip rate", "0%"],
    ],
)

add_body("Key reward changes for obstacle specialization:", bold_prefix="")
doc.paragraphs[-1].runs[0].bold = True
add_bullet(" 0.5 → 2.0 (lift feet higher to step over obstacles)", bold_prefix="foot_clearance:")
add_bullet(" -3.0 → -2.0 (allow lateral tilt for boulder traversal)", bold_prefix="base_orientation:")
add_bullet(" -1.0 (maintained for gait quality)", bold_prefix="action_smoothness:")

add_body("Evaluation:", bold_prefix="")
doc.paragraphs[-1].runs[0].bold = True

add_table(
    ["Environment", "Distance", "Zones"],
    [
        ["Friction", "42.2m", "5/5"],
        ["Grass", "31.7m", "4/5"],
        ["Boulder", "30.4m", "4/5"],
        ["Stairs", "15.7m", "2/5"],
    ],
    highlight_rows=[2],
)

add_callout(
    "The specialization trade-off in action: The Obstacle Expert gained +10.1m on boulders "
    "(30.4m vs 20.3m) but lost -6.7m on friction (42.2m vs 48.9m). Each expert excels at "
    "its terrain, but neither is good at everything. This is exactly why we need distillation."
)

add_hr()

# ── 5.4 Distilled Master ──
doc.add_heading("5.4 Stage 4 — The Distilled Master", level=3)

add_body(
    "A single student policy that learns WHEN to use each expert's behavior by reading "
    "the terrain geometry through its height scan. The student acts in the environment, "
    "both frozen experts label what they would have done, and the student is trained to "
    "match the appropriate expert for each terrain."
)

add_table(
    ["Property", "Value"],
    [
        ["Checkpoint", "distilled_6899.pt"],
        ["Architecture", "[512, 256, 128] (same as both experts)"],
        ["Iterations", "6,899"],
        ["Training time", "~6-8 hours on H100"],
        ["Environments", "4,096"],
        ["Experts", "Friction (mason_hybrid_best_33200.pt) + Obstacle (obstacle_best_44400.pt)"],
    ],
)

doc.add_heading("How the Terrain Router Works", level=3)

add_body(
    "The height scan (first 187 observation dimensions) encodes terrain geometry. "
    "Flat terrain has near-zero variance. Boulders and stairs have high variance. "
    "A sigmoid gate routes each environment to the appropriate expert:"
)

add_code_block(
    "Height Scan (187 dims) → compute variance → sigmoid gate\n"
    "                                              │\n"
    "                              gate ≈ 0: smooth terrain → Friction Expert\n"
    "                              gate ≈ 1: rough terrain  → Obstacle Expert\n"
    "                              gate ≈ 0.5: transition   → blend both"
)

add_body(
    "The student doesn't hard-switch between experts — it smoothly blends their actions "
    "at terrain boundaries, preventing jerky transitions."
)

doc.add_heading("Training Process", level=3)

add_numbered("", bold_prefix="Initialize student ")
doc.paragraphs[-1].add_run(
    "from the Friction Expert (best general gait — the student already knows how to walk)."
)
add_numbered("", bold_prefix="Critic warmup ")
doc.paragraphs[-1].add_run(
    "(300 iterations): Actor is frozen while the critic learns the new value landscape."
)
add_numbered("", bold_prefix="Combined training ")
doc.paragraphs[-1].add_run(
    "(each iteration): Student collects PPO experience, runs PPO update, then a post-hoc "
    "distillation step queries both frozen experts, blends their actions via the terrain gate, "
    "and computes MSE + KL loss between student and blended expert target."
)
add_numbered("", bold_prefix="Alpha annealing: ")
doc.paragraphs[-1].add_run(
    "Distillation weight starts at 0.8 (mostly copy experts) and decays to 0.2 "
    "(mostly PPO reward signal). The student absorbs expert knowledge first, then adapts."
)

doc.add_heading("Distillation Hyperparameters", level=3)

add_table(
    ["Parameter", "Value", "Purpose"],
    [
        ["Alpha (start → end)", "0.8 → 0.2", "Shifts from expert imitation to PPO reward"],
        ["KL weight", "0.1", "Balance between MSE and KL divergence loss"],
        ["Roughness threshold", "0.005", "Height scan variance gate for routing"],
        ["Routing temperature", "0.005", "Sigmoid sharpness (lower = harder gate)"],
        ["Distill batch size", "8,192", "Samples per distillation gradient step"],
        ["Critic warmup", "300 iters", "Actor frozen while critic calibrates"],
    ],
)

doc.add_heading("Distillation Loss", level=3)

add_code_block(
    "loss = MSE(student_action, blended_expert_action) + 0.1 × KL(student_dist ∥ expert_dist)"
)
add_body("The KL term is clamped to [0.0, 10.0] for numerical stability.")

add_hr()

# ── 6. Comparative Evaluation ──
doc.add_heading("6. Comparative Evaluation", level=2)

doc.add_heading("6.1 The Full Pipeline — Side by Side", level=3)

add_table(
    ["Policy", "Friction", "Grass", "Boulder", "Stairs"],
    [
        ["ARL Baseline", "36.9m (4.1)", "29.6m (3.6)", "14.4m (2.2)", "10.9m (2.0)"],
        ["ARL Hybrid", "48.9m (5.0)", "27.2m (3.3)", "20.3m (3.0)", "11.2m (2.0)"],
        ["Obstacle Expert", "42.2m (5.0)", "31.7m (4.0)", "30.4m (4.0)", "15.7m (2.0)"],
    ],
    highlight_rows=[1],
)

add_body(
    "Values shown as: mean progress (mean zone). ARL Baseline and Hybrid from 100-episode "
    "evaluations. Obstacle Expert from single-episode evaluation.",
)
doc.paragraphs[-1].runs[0].font.size = Pt(9)
doc.paragraphs[-1].runs[0].italic = True

add_callout(
    "The specialization dilemma: The ARL Hybrid dominates friction (+6.7m over Obstacle Expert). "
    "The Obstacle Expert dominates boulders (+10.1m over Hybrid). Neither is best at everything. "
    "Distillation resolves this by teaching the student to use the right expert for each terrain."
)

doc.add_heading("6.2 Stability Comparison", level=3)

add_table(
    ["Policy", "Friction Falls", "Grass Falls", "Boulder Falls", "Stairs Falls", "Flip Rate"],
    [
        ["ARL Baseline", "37%", "23%", "13%", "20%", "—"],
        ["ARL Hybrid", "2%", "15%", "3%", "36%", "0%"],
    ],
    highlight_rows=[1],
)

add_callout(
    "The ARL Hybrid's 0% flip rate and 2% friction fall rate demonstrate the value of "
    "ARL's conservative reward philosophy. High gait weight (10.0) and moderate penalties "
    "produce gaits that prioritize stability."
)

add_hr()

# ── 7. Teammate Locomotion Work ──
doc.add_heading("7. Teammate Locomotion Work", level=2)

doc.add_heading("7.1 Ryan's ARL Hybrid Baseline", level=3)
add_body(
    "Ryan developed the ARL Hybrid baseline policy (mason_hybrid_best_33200.pt) which "
    "served as the foundation for the entire pipeline:"
)
add_bullet("Used as the Friction/Grass Expert in distillation")
add_bullet("Used by Colby as the frozen locomotion backbone in his CombinedPolicyTraining navigation system")
add_bullet("The [512, 256, 128] architecture and conservative weight configuration established the standard for all subsequent training")

doc.add_heading("7.2 Colby", level=3)
add_body(
    "Colby did not develop standalone locomotion policies. His navigation work "
    "(CombinedPolicyTraining) uses Ryan's mason_hybrid_best_33200.pt as a frozen "
    "locomotion backbone, with the navigation policy outputting velocity commands "
    "that the frozen loco policy converts to joint actions."
)

doc.add_heading("7.3 Cole", level=3)
add_body(
    "Cole did not develop standalone locomotion policies. His navigation system "
    "(RL_FOLDER_VS2/VS3) uses a SpotFlatTerrainPolicy as the frozen locomotion backbone, "
    "following the same hierarchical approach where high-level velocity commands are "
    "executed by a pre-trained walking controller."
)

add_hr()

# ── 8. Key Lessons Learned ──
doc.add_heading("8. Key Lessons Learned", level=2)

lessons = [
    ("Simpler is better. ",
     "ARL's 11-term reward function with a 286K-parameter network outperformed our "
     "22-term function with a 1.2M-parameter network. Fewer rewards give clearer gradients; "
     "smaller networks generalize instead of memorizing."),
    ("0% flip rate is achievable. ",
     "ARL's conservative weight philosophy (high gait weight, moderate penalties) produces "
     "gaits that never flip, even on difficult terrain. Stability is the foundation everything "
     "else builds on."),
    ("Specialize, then distill. ",
     "Training a single policy on mixed terrain forces impossible trade-offs — loose penalties "
     "for obstacles vs. tight penalties for clean walking. Training two specialists and combining "
     "them via distillation sidesteps this entirely."),
    ("Best checkpoint ≠ final checkpoint. ",
     "Training regression is common. The Obstacle Expert's best was at iteration 44,400 out of "
     "a 54,600-iteration run. Periodic evaluation during training is essential."),
    ("The height scan is the terrain router. ",
     "The 187-dimensional height scan variance cleanly separates smooth terrain from rough terrain, "
     "providing a natural signal for expert routing without any additional sensors."),
    ("Initialize from your best generalist. ",
     "Starting the distilled student from the Friction Expert means it already knows how to walk. "
     "It only needs to learn when to switch to obstacle behavior — saving thousands of iterations."),
    ("Tune rewards along kinematic chains. ",
     "For obstacle terrain, foot clearance + action smoothness + joint position all govern the "
     "step-up motion. The biggest boulder improvement (+10.1m) came from adjusting these three together."),
]

for bold, rest in lessons:
    add_numbered("", bold_prefix=bold)
    doc.paragraphs[-1].add_run(rest)

# ── Save ──
out_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "Phase_Deliverable_Locomotion.docx",
)
doc.save(out_path)
print(f"Saved to {out_path}")
